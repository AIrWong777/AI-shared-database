"""
文本提取工具函数
支持从PDF、DOCX和HTML文件中提取文本和元数据
"""

import os
from pathlib import Path
from typing import Optional, Dict, Any
import logging
import re

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def clean_text(text: str) -> str:
    """
    清理提取的文本
    
    Args:
        text: 原始文本
        
    Returns:
        str: 清理后的文本
    """
    if not text:
        return ""
    
    # 1. 替换多个空白字符为单个空格
    text = re.sub(r'\s+', ' ', text)
    
    # 2. 清理特殊字符，但保留基本标点
    text = re.sub(r'[^\w\s,.!?;:()\-\'\"，。！？；：（）]', '', text)
    
    # 3. 清理可能的分页符号
    text = re.sub(r'-\s+', '', text)
    
    # 4. 清理连续的标点符号
    text = re.sub(r'[,.!?;:]{2,}', '.', text)
    
    # 5. 确保句子之间有适当的空格
    text = re.sub(r'([.!?])\s*([A-Z])', r'\1 \2', text)
    
    return text.strip()

def extract_pdf_text(file_path: str) -> Optional[str]:
    """
    使用PyPDF2从PDF文件中提取文本
    
    Args:
        file_path: PDF文件路径
        
    Returns:
        Optional[str]: 提取的文本内容，失败时返回None
    """
    try:
        import PyPDF2
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # 提取所有页面的文本
            text_parts = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_parts.append(page.extract_text())
            
            # 合并所有页面的文本
            text = "\n".join(text_parts)
            
            # 清理文本
            text = clean_text(text)
            
            if text.strip():
                logger.info(f"成功从PDF提取文本，长度: {len(text)} 字符")
                return text
            else:
                logger.warning(f"PDF文件 {file_path} 中没有可提取的文本")
                return None
                
    except ImportError:
        logger.error("PyPDF2库未安装，无法提取PDF文本")
        return None
    except Exception as e:
        logger.error(f"提取PDF文本失败: {e}")
        return None

def extract_docx_text(file_path: str) -> Optional[str]:
    """
    从Word文档中提取文本
    
    Args:
        file_path: DOCX文件路径
        
    Returns:
        Optional[str]: 提取的文本内容，失败时返回None
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        # 提取所有段落的文本
        paragraphs_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs_text.append(paragraph.text)
        
        # 提取所有表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs_text.append(cell.text)
        
        # 合并所有文本
        text = "\n".join(paragraphs_text)
        
        # 清理文本
        text = clean_text(text)
        
        if text.strip():
            logger.info(f"成功从DOCX提取文本，长度: {len(text)} 字符")
            return text
        else:
            logger.warning(f"DOCX文件 {file_path} 中没有可提取的文本")
            return None
            
    except ImportError:
        logger.error("python-docx库未安装，无法提取DOCX文本")
        return None
    except Exception as e:
        logger.error(f"提取DOCX文本失败: {e}")
        return None

def extract_html_text(file_path: str) -> Optional[str]:
    """
    从HTML文件中提取文本
    
    Args:
        file_path: HTML文件路径
        
    Returns:
        Optional[str]: 提取的文本内容，失败时返回None
    """
    try:
        from bs4 import BeautifulSoup
        
        with open(file_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file, 'lxml')
            
            # 移除script和style元素
            for script in soup(["script", "style"]):
                script.decompose()
            
            # 获取文本
            text = soup.get_text()
            
            # 清理文本
            text = clean_text(text)
            
            if text.strip():
                logger.info(f"成功从HTML提取文本，长度: {len(text)} 字符")
                return text
            else:
                logger.warning(f"HTML文件 {file_path} 中没有可提取的文本")
                return None
                
    except ImportError:
        logger.error("beautifulsoup4或lxml库未安装，无法提取HTML文本")
        return None
    except Exception as e:
        logger.error(f"提取HTML文本失败: {e}")
        return None

def extract_title_from_text(text: str, max_length: int = 50) -> str:
    """
    从文本中提取标题（改进版）
    
    Args:
        text: 文本内容
        max_length: 标题最大长度
        
    Returns:
        str: 提取的标题
    """
    if not text or not text.strip():
        return "未知标题"
    
    # 清理文本
    text = text.strip()
    
    # 按行分割
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    if not lines:
        return "未知标题"
    
    # 选择第一个非空行作为候选标题
    candidate_title = lines[0]
    
    # 如果第一行太短，且第二行存在，考虑组合前两行
    if len(lines) > 1 and len(candidate_title) < 10:
        combined_title = f"{candidate_title} - {lines[1]}"
        if len(combined_title) <= max_length:
            candidate_title = combined_title
    
    # 限制标题长度
    if len(candidate_title) > max_length:
        candidate_title = candidate_title[:max_length] + "..."
    
    return candidate_title

def extract_title_from_filename(filename: str) -> str:
    """
    从文件名中提取标题（去除扩展名）
    
    Args:
        filename: 文件名
        
    Returns:
        str: 提取的标题
    """
    return Path(filename).stem

def extract_metadata_from_file(file_path: str, original_filename: str) -> Dict[str, Any]:
    """
    从文件中提取元数据（标题等）
    
    Args:
        file_path: 文件路径
        original_filename: 原始文件名
        
    Returns:
        Dict[str, Any]: 包含提取元数据的字典
    """
    file_ext = Path(file_path).suffix.lower()
    
    metadata = {
        "title": extract_title_from_filename(original_filename),
        "extracted_text": None,
        "extraction_success": False,
        "file_type": file_ext[1:] if file_ext else "unknown"  # 去掉点号
    }
    
    # 根据文件类型选择相应的提取函数
    if file_ext == '.pdf':
        extracted_text = extract_pdf_text(file_path)
    elif file_ext == '.docx':
        extracted_text = extract_docx_text(file_path)
    elif file_ext in ['.html', '.htm']:
        extracted_text = extract_html_text(file_path)
    else:
        logger.warning(f"不支持的文件类型: {file_ext}")
        return metadata
    
    if extracted_text:
        # 从提取的文本中获取更好的标题
        title_from_text = extract_title_from_text(extracted_text)
        if title_from_text and title_from_text != "未知标题":
            metadata["title"] = title_from_text
        
        metadata["extracted_text"] = extracted_text
        metadata["extraction_success"] = True
        metadata["text_length"] = len(extracted_text)
        
        logger.info(f"成功提取文件元数据: {metadata['title']}")
    else:
        logger.warning(f"文本提取失败，使用文件名作为标题")
    
    return metadata

def is_text_extractable(file_path: str) -> bool:
    """
    检查文件是否支持文本提取
    
    Args:
        file_path: 文件路径
        
    Returns:
        bool: 是否支持文本提取
    """
    file_ext = Path(file_path).suffix.lower()
    # 支持PDF、DOCX和HTML
    return file_ext in ['.pdf', '.docx', '.html', '.htm']